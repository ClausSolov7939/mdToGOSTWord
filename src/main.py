import json

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, Mm


setupFile = open("ОтДуксина.json", encoding='utf-8')
setupData = json.load(setupFile)

document = Document()

section = document.sections[0]

pageData = setupData["Page Setup"]

# section.page_height = Mm(297)
# section.page_width = Mm(210)
section.left_margin = Mm(pageData["left_margin"])
section.right_margin = Mm(pageData["right_margin"])
section.top_margin = Mm(pageData["top_margin"])
section.bottom_margin = Mm(pageData["bottom_margin"])

styles = document.styles
styleData = setupData["Styles"]

for i in styles:
    styles[i.name].delete()

for newStyle in styleData.keys():
    currentStyle = styles.add_style(newStyle, WD_STYLE_TYPE.PARAGRAPH)
    # print(currentStyle)
    # -------- font setting -----------
    font_format = currentStyle.font
    font_style = styleData[newStyle]["font"]
    font_format.name = font_style["name"]
    font_format.size = Pt(font_style["size"])
    font_format.bold = font_style["bold"]
    font_format.italic = font_style["italic"]
    font_format.underline = font_style["underline"]
    font_format.all_caps = font_style["all_caps"]
    # -------- paragraph setting -----------
    paragraph_format = currentStyle.paragraph_format
    paragraph_style = styleData[newStyle]["paragraph"]
    # 0 - left, 1 - center, 2 - right, 3 - justify

    alignment = {"left": 0, "center": 1, "right": 2, "justify": 3}

    paragraph_format.alignment = alignment[paragraph_style["alignment"]]
    print(paragraph_style["alignment"], paragraph_format.alignment)
    paragraph_format.first_line_indent = Mm(
        paragraph_style["first_line_indent"])
    paragraph_format.left_indent = Mm(paragraph_style["left_indent"])
    paragraph_format.right_indent = Mm(paragraph_style["right_indent"])
    paragraph_format.space_before = Mm(paragraph_style["space_before"])
    paragraph_format.space_after = Mm(paragraph_style["space_after"])
    if paragraph_style["line_spacing"] == 1.5:
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    elif paragraph_style["line_spacing"] == 1:
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    # Запрет висящих строк
    paragraph_format.widow_control = paragraph_style["widow_control"]
    # Не отрывать от следующего
    paragraph_format.keep_with_next = paragraph_style["keep_with_next"]
    # Не разрывать абзац
    paragraph_format.keep_together = paragraph_style["keep_together"]
    # с новой страницы
    paragraph_format.page_break_before = paragraph_style["page_break_before"]

# print(document.styles["Заголовок 1"])
# print(document.styles["Заголовок 1"].paragraph_format.page_break_before)

for i in styles:
    print(i.name)

fin = open("../input/untitled.md", "r", encoding="utf-8")
for line in fin:
#    print("[ ",line," ]")
    if line!="\n":
        # Заголовки
        if "#" in line.split(" ")[0]:
            head=line.replace("#","").replace("\n","")
            level = len(line.split(" ")[0])
            document.add_paragraph(head[1:], style=f"Head {level}")
        elif line[0]=="!":
            imagePath = "../input/"+line.replace("![[","").replace("]]\n","")
            print(imagePath)
#            doc.add_paragraph("Picture",style="Normal")
            p = document.add_picture(imagePath, width=Mm(153))
#            p.style.paragraph_format.alignment = 1
            last_paragraph = document.paragraphs[-1]
            last_paragraph.style = "Picture"
            document.add_paragraph(fin.readline(), style="PictureName")
        else:
            document.add_paragraph(line.replace("\n",""),style="Main Text")
    else:
        print("empty")

document.save('../output/demo.docx')
